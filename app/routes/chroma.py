from fastapi import APIRouter, status, Depends, HTTPException, UploadFile, File, Form
from app.services.chroma_service import ChromaService
from app.dependencies.auth import req_admin_role  # Admin-only dependency
from docx import Document
import os

router = APIRouter()
chroma_service = ChromaService()

### ✅ ADD Pokémon Context (with .docx upload)
@router.post("/add-context", status_code=status.HTTP_201_CREATED, dependencies=[Depends(req_admin_role)])
async def add_pokemon_context(
    pokemon_id: str = Form(...),
    pokemon_name: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Admin-only endpoint to add Pokémon context by uploading a Word document (.docx).
    """
    try:
        if not file.filename.endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx files are supported.")

        # Save uploaded file temporarily
        temp_path = f"temp_{file.filename}"
        contents = await file.read()
        with open(temp_path, "wb") as f:
            f.write(contents)

        # Parse Word file content
        doc = Document(temp_path)
        parsed_text = ""
        for para in doc.paragraphs:
            parsed_text += para.text.strip() + "\n\n"

        # Save to ChromaDB by pokemon_id
        chroma_service.add_pokemon(pokemon_id, pokemon_name, parsed_text)

        # Clean up temp file
        os.remove(temp_path)

        return {"message": f"Context for Pokémon '{pokemon_name}' (ID: {pokemon_id}) added successfully."}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


### ✅ SEARCH Pokémon Context by ID
@router.get("/search-context/{pokemon_id}", dependencies=[Depends(req_admin_role)])
async def search_pokemon_context(pokemon_id: str):
    try:
        results = chroma_service.search_pokemon(pokemon_id)
        if not results:
            raise HTTPException(status_code=404, detail="No context found.")
        return {"context": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


### ✅ DELETE Pokémon Context by ID
@router.delete("/delete-context/{pokemon_id}", dependencies=[Depends(req_admin_role)])
async def delete_pokemon_context(pokemon_id: str):
    try:
        deleted = chroma_service.delete_pokemon(pokemon_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="No context found to delete.")
        return {"message": f"Deleted context for Pokémon ID: {pokemon_id}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


### ✅ UPDATE Pokémon Context (Replace existing)
@router.put("/update-context", dependencies=[Depends(req_admin_role)])
async def update_pokemon_context(
    pokemon_id: str = Form(...),
    pokemon_name: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Admin-only endpoint to update Pokémon context.
    """
    try:
        # Delete existing context (if exists)
        chroma_service.delete_pokemon(pokemon_id)

        # Reuse add logic:
        if not file.filename.endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx files are supported.")

        temp_path = f"temp_{file.filename}"
        contents = await file.read()
        with open(temp_path, "wb") as f:
            f.write(contents)

        doc = Document(temp_path)
        parsed_text = ""
        for para in doc.paragraphs:
            parsed_text += para.text.strip() + "\n\n"

        chroma_service.add_pokemon(pokemon_id, pokemon_name, parsed_text)
        os.remove(temp_path)

        return {"message": f"Context for Pokémon '{pokemon_name}' (ID: {pokemon_id}) updated successfully."}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
